import os
from typing import List, Dict, Any, Optional, Tuple
import pandas as pd
import PyPDF2
from docx import Document
import re
import config

class DocumentProcessor:
    """Handles document ingestion and text extraction"""
    
    def __init__(self):
        self.supported_formats = config.ALLOWED_EXTENSIONS
    
    def process_file(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """
        Process uploaded file and extract text with metadata
        
        Returns:
            Dict containing text, metadata, and source information
        """
        file_extension = os.path.splitext(file_name)[1].lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(file_path, file_name)
        elif file_extension in ['.xlsx', '.xls']:
            return self._process_excel(file_path, file_name)
        elif file_extension == '.docx':
            return self._process_docx(file_path, file_name)
        elif file_extension == '.txt':
            return self._process_txt(file_path, file_name)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract text from PDF with enhanced processing and page-level metadata"""
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                # Enhanced text cleaning
                if text.strip():
                    # Clean up common PDF extraction artifacts
                    text = self._clean_pdf_text(text)
                    
                    # Only add if text meets minimum length requirement
                    if len(text.strip()) >= config.MIN_CHUNK_SIZE:
                        documents.append({
                            'text': text,
                            'metadata': {
                                'source': file_name,
                                'page': page_num + 1,
                                'type': 'pdf',
                                'total_pages': len(pdf_reader.pages),
                                'char_count': len(text),
                                'word_count': len(text.split())
                            }
                        })
        
        return {'documents': documents, 'source_type': 'pdf'}
    
    def _clean_pdf_text(self, text: str) -> str:
        """Clean PDF extracted text to improve quality"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common PDF extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between words
        text = re.sub(r'(\.)([A-Z])', r'\1 \2', text)     # Add space after periods
        text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)  # Add space between letters and numbers
        text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)  # Add space between numbers and letters
        
        # Remove page headers/footers patterns (common artifacts)
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            # Skip very short lines that are likely headers/footers
            if len(line) < 3:
                continue
            # Skip lines that are just page numbers or dates
            if re.match(r'^\d+$', line) or re.match(r'^\d{1,2}[/-]\d{1,2}[/-]\d{2,4}$', line):
                continue
            cleaned_lines.append(line)
        
        return ' '.join(cleaned_lines).strip()
    
    def _process_excel(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract text from Excel with sheet and row metadata"""
        documents = []
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Convert each row to text
            for idx, row in df.iterrows():
                # Create readable text from row data
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                
                text = " | ".join(row_text)
                
                if text.strip():
                    documents.append({
                        'text': text,
                        'metadata': {
                            'source': file_name,
                            'sheet': sheet_name,
                            'row': idx + 2,  # +2 because pandas is 0-indexed and Excel starts at 1, plus header
                            'type': 'excel',
                            'columns': list(df.columns)
                        }
                    })
        
        return {'documents': documents, 'source_type': 'excel'}
    
    def _process_docx(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract text from Word document with paragraph metadata"""
        documents = []
        
        doc = Document(file_path)
        
        for para_num, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                documents.append({
                    'text': text,
                    'metadata': {
                        'source': file_name,
                        'paragraph': para_num + 1,
                        'type': 'docx',
                        'total_paragraphs': len(doc.paragraphs)
                    }
                })
        
        return {'documents': documents, 'source_type': 'docx'}
    
    def _process_txt(self, file_path: str, file_name: str) -> Dict[str, Any]:
        """Extract text from plain text file"""
        documents = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
            # Split by paragraphs (double newlines) or lines if no paragraphs
            paragraphs = content.split('\n\n')
            if len(paragraphs) == 1:
                paragraphs = content.split('\n')
            
            for para_num, paragraph in enumerate(paragraphs):
                text = paragraph.strip()
                if text:
                    documents.append({
                        'text': text,
                        'metadata': {
                            'source': file_name,
                            'paragraph': para_num + 1,
                            'type': 'txt',
                            'total_paragraphs': len(paragraphs)
                        }
                    })
        
        return {'documents': documents, 'source_type': 'txt'}

class TextChunker:
    """Handles text chunking for optimal embedding generation"""
    
    def __init__(self, chunk_size: int = config.CHUNK_SIZE, chunk_overlap: int = config.CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_documents(self, documents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Split documents into smaller chunks with enhanced boundary detection
        """
        chunked_documents = []
        
        for doc in documents:
            text = doc['text']
            metadata = doc['metadata']
            
            # If text is smaller than chunk size, keep as is (but check minimum size)
            if len(text) <= self.chunk_size:
                if len(text.strip()) >= config.MIN_CHUNK_SIZE:
                    chunked_documents.append(doc)
                continue
            
            # Split into chunks with improved boundary detection
            chunks = self._split_text_enhanced(text)
            
            for chunk_num, chunk in enumerate(chunks):
                # Only add chunks that meet minimum size requirement
                if len(chunk.strip()) >= config.MIN_CHUNK_SIZE:
                    chunk_metadata = metadata.copy()
                    chunk_metadata['chunk'] = chunk_num + 1
                    chunk_metadata['total_chunks'] = len(chunks)
                    chunk_metadata['chunk_char_count'] = len(chunk)
                    chunk_metadata['chunk_word_count'] = len(chunk.split())
                    
                    chunked_documents.append({
                        'text': chunk,
                        'metadata': chunk_metadata
                    })
        
        return chunked_documents
    
    def _split_text_enhanced(self, text: str) -> List[str]:
        """Enhanced text splitting with better boundary detection"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            # If this isn't the last chunk, try to find optimal boundary
            if end < len(text):
                # Look for paragraph boundaries first (double newline)
                paragraph_end = text.rfind('\n\n', start, end)
                if paragraph_end > start + self.chunk_size // 2:
                    end = paragraph_end + 2
                else:
                    # Look for sentence endings
                    for i in range(end, max(start, end - 200), -1):
                        if text[i] in '.!?':
                            # Make sure it's actually end of sentence
                            if i + 1 < len(text) and text[i + 1] in ' \n':
                                end = i + 1
                                break
                    else:
                        # Look for word boundaries
                        for i in range(end, max(start, end - 100), -1):
                            if text[i] in ' \n\t':
                                end = i
                                break
            
            chunk = text[start:end].strip()
            if chunk and len(chunk) >= config.MIN_CHUNK_SIZE:
                chunks.append(chunk)
            
            # Move start position with overlap, but ensure progress
            new_start = max(end - self.chunk_overlap, start + 1)
            if new_start >= end:
                break
            start = new_start
        
        return chunks

    def _split_text(self, text: str) -> List[str]:
        """Original text splitting method (fallback)"""
        return self._split_text_enhanced(text)